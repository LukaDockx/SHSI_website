
file_path = 'data2.csv' 
docxprintdoc = 'Attendance_sheet_t3.docx' 
database_path = 'new_database_s2.csv'
yesterday_housing_attendance = 'data_yesterday.csv'
facinfopath = 'phonenr_s2.csv'

# Turning False will not generate .docx files for people where the faculty hasn't provided presence/absence status.
print_notspecified = True
#

'''
#-------------------------------------------------------------#
#-------------------------------------------------------------#
!!! ALL .CSV DOCUMENTS SHOULD BE UTF-8 ENCODED !!!

!!! MAKE SURE THE .CSV FILES DON'T HAVE EMPTY LINES AFTER THEM !!!

Make sure NONE of the .csv and .docx files are open in a different program.

When updating the database (stored in database_path), remove session 2 people.

1. The current attendance data should be stored under the 'file_path' variable.

2. The attendance report from yesterday should be stored in the 'yesterday_housing_attendance' variable 
   this is used to check whether the student was in their dorm yesterday.

3. You can add / change phone numbers by adjusting the third column in the file linked to the 'facinfopath' variable.

4. The 'docxprintdoc' variable determines the name of the final .docx file that will be generated.
   Make sure this file is not open in a different window or python will not have permission to write to it.
   
5. The print_notspecified variable will remove people for which no attendance data is available, use with caution.

#-------------------------------------------------------------#
#-------------------------------------------------------------#

'''


import pandas as pd
Houses = ['Carter House 1',
'School District Waitlist',
'Carter House 2',
'Carter House 3',
'Casa Jackson 1',
'Casa Jackson 2',
'Casa Jackson 3',
'Casa Werner 1',
'Casa Werner 2',
'Casa Werner 3',
'Eiselen House 1',
'Eiselen House 2',
'Eiselen House 3',
'Farley House 1',
'Farley House 2',
'Farley House 3',
'Grace Covell Hall 1',
'Grace Covell Hall 2',
'Grace Covell Hall 3',
'Jessie Ballantyne Hall 1',
'Jessie Ballantyne Hall 2',
'Jessie Ballantyne Hall 3',
'John Ballantyne Hall 1',
'John Ballantyne Hall 2',
'John Ballantyne Hall 3',
'Price House 1',
'Price House 2',
'Price House 3',
'Ritter House 1',
'Ritter House 2',
'Ritter House 3',
'Wemyss Hall 1',
'School District Waitlist housing',
# NEW ENTRIES: ---
# 'Grace Covell Hall',
# 'Eiselen House',
# 'Jessie Ballantyne Hall',
# 'John Ballantyne Hall',
# 'Wemyss Hall',
# 'Wemyss House',
# 'Ritter House',
# 'Casa Werner',
# 'Casa Jackson',
# 'Carter House',
# 'Farley House',
# 'Calaveras Hall',
# 'Calaveras Hall 3',
# 'Price House',
# 'Calaveras Hall 2',
# done ---
'Wemyss Hall 2',
'Wemyss Hall 3']

Activities = ["2D Art Studio",
"School District Waitlist",
"3D Modeling and Printing",
"Arte Tridimensional y Animación",
"Baseball Skills and Tactics",
"Basketball Skills Academy",
"Be a Lawyer in Court",
"Beach Volleyball",
"Become the Best Teacher or Coach",
"Building Electrical Circuits and Computers",
"Building Smart Devices",
"Ceramics Studio",
"Choir",
"Coding",
"Competitive Debate",
"Computer Drawing and Drafting",
"Creative Writing and Storytelling",
"Data Science",
"Design and Synthesis of Anticancer Drugs",
"Drone Flying",
"eSports",
"Esport",
"Exploring the Human Body",
"F1tenth Racing",
"Finding a Cure for Cancer",
"Future Dentists",
"Future Pharmacists",
"Instrumental Music",
"Investing in Stocks",
"Jazz",
"Leadership for Women",
"Musical Theatre Workshop & Showcase",
"Nursing: Caring Hands, Healing Hearts",
"Piano",
"Plant Biodiversity",
"Soccer Skills and Tactics - Week 2/4",
"Soccer Skills and Tactics - Week 1/3",
"Soccer Skills and Tactics",
"Softball",
"Sports Analytics",
"Taking Control of Your Financial Future",
"Tennis Skills and Tactics",
"The Science and Practice of Fitness",
"Virtual Reality and Immersive Design",
"Volleyball Skills and Tactics",
"Water Polo Skills and Tactics"]


print('\n (Note) All .csv documents should be UTF-8 encoded!\n')


def garbage_remover(cell):
    if isinstance(cell, str):
        return ' '.join(cell.split())
    else:
        return cell

def remove_empty_rows(file_path):
    """
    1) If the very first cell of the CSV equals
       'Live attendance - Pacific Summer High School Institute',
       drop the first two lines before anything else.
    2) Remove any data‐rows that consist ONLY of commas (and/or whitespace).
    3) Overwrite the original file with the cleaned content.
    """

    # Read all lines from the file
    with open(file_path, 'r', encoding='utf-8') as f:
        raw_lines = f.readlines()

    # If there's nothing in the file, just return
    if not raw_lines:
        return

    # 1) Check the first cell of the very first line
    first_cell = raw_lines[0].split(',')[0].strip()
    if "Live" in first_cell:
        # Drop the first two lines entirely
        lines = raw_lines[4:]
    else:
        lines = raw_lines.copy()

    # If we removed all lines by skipping, nothing to do
    if not lines:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.writelines([]) 
        return

    # The new header is whatever is now line 0 after skipping
    header = lines[0]
    data_lines = lines[1:]

    cleaned_data = []
    for line in data_lines:
        # 2) If a line is blank (just whitespace) → skip
        if line.strip() == "":
            continue

        # Remove spaces/tabs/newlines, then remove commas:
        # If what's left is empty, it was a “only commas/whitespace” row → skip
        content_only = "".join(ch for ch in line if ch not in ", \t\r\n")
        if content_only == "":
            continue

        # Otherwise, keep it
        cleaned_data.append(line)

    # Recombine header + cleaned data
    new_lines = [header] + cleaned_data

    # 3) Overwrite the CSV with these cleaned lines
    with open(file_path, 'w', encoding='utf-8') as f:
        f.writelines(new_lines)
    
# REMOVE EMPTY ROWS
print("(WARNING) IF YOU GET AN (permission) ERROR IMMEDIATLY AFTER THIS MESSAGE, YOU NEED TO CLOSE ALL THE EXCEL/CSV FILES, THE PROGRAM DOESN\'T HAVE PERMISION TO READ IT")
remove_empty_rows(file_path)
remove_empty_rows(database_path)
remove_empty_rows(yesterday_housing_attendance)
remove_empty_rows(facinfopath)
#
    
def remove_in_house(df):
    df = df[~df['Program'].isin(Houses)]
    invalid_programs = df[~df['Program'].isin(Houses + Activities)]
    if not invalid_programs.empty:
        print('invalid programs:')
        print(invalid_programs)
        print(invalid_programs['Program'])
        raise ValueError("Error: Some entries in 'Program' are neither in Houses nor Activities. Please add these to housing or activities.")
    return df

#Remove anyone that is checked out!
def get_invalid_participants(df):
    invalid_participants = df[(df['Status'].isin(['Checked out', 'Not checked in'])) & (df['Attendance Status'] == 'Present')]['Participant']
    return invalid_participants
    
def remove_not_checkedin(df):
    df = df[df['Status'] != 'Checked out']
    df = df[df['Status'] != 'Not checked in']
    return df

# Check for empty entries in the 'Attendance Status' column
def check_empty_entries(df):
    empty_entries = df[df['Attendance Status'].isna()]
    
    all_notsubprog = {}
    for person in empty_entries.to_numpy():
        if person[4] in all_notsubprog:
            all_notsubprog[person[4]] += 1
        else:
            all_notsubprog[person[4]] = 1


    if not empty_entries.empty:
        print('# ------------ (WARNING) Some Faculty hasn\'t provided their attendance status: ------------ #')
        print(all_notsubprog)
        print(empty_entries.to_numpy().T[0])
        print('# ------------------------------------------------------------------------------------------ #\n\n')
    else:
        print("(LOGGER) All Faculty has provided their attendance status.\n")
    return empty_entries

def remove_empty_entries(df):
    empty_entries = df[df['Attendance Status'].isna()]
    filled_entries = non_empty_entries = df[df['Attendance Status'].notna()]
    
    all_notsubprog = {}
    
    for person in empty_entries.to_numpy():
        if person[4] in all_notsubprog:
            all_notsubprog[person[4]] += 1
        else:
            all_notsubprog[person[4]] = 1
    return empty_entries, filled_entries

def check_notpresent(df):
    # Check for 'Attendance Status' values that are not 'Present' or 'Late'
    invalid_statuses = df[~df['Attendance Status'].isin(['Present', 'Late'])]

    if not invalid_statuses.empty:
        print("# ------ Participants with status other than \'Present\' or \'Late\' ------ #")
        for persona in invalid_statuses['Participant']:
            print(persona)
            
        print('\nPacific ID\'s:')
        for persona in invalid_statuses['Participant external ID']:
            print(persona)        
        
        print("# ---------------------------------------------------------------------- #\n\n")
    else:
        print("(LOGGER) All participants have attendance statuses of either 'Present' or 'Late'.")
    return invalid_statuses

# Load the CSV file into a DataFrame
df = pd.read_csv(file_path, encoding = "utf-8")
df = df.map(garbage_remover)


import re

# Pattern: comma + space + 3-letter month + rest-of-line
date_pattern = r',\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b.*$'

df['Program'] = (
    df['Program']
      .astype(str)                        # ensure string dtype
      .str.replace(date_pattern, '',      # drop “, fJun 2, 2025(…”
                   regex=True)
      .str.strip()                        # clean up any trailing spaces
)


df = remove_in_house(df)
present_notcheckedin = get_invalid_participants(df)
print('# ------------ present but not checked in ------------ #')
for persona in present_notcheckedin:
    print(persona)
#print(present_notcheckedin)
print('# ---------------------------------------------------- #\n\n')

df = remove_not_checkedin(df)

empty_entries = check_empty_entries(df)

if print_notspecified == False:
    print('# --- (WARNING) not generating a report for students with missing absence/presence data. --- #')
    print('# --- Set The \'print_notspecified\' Variable to TRUE if you want a report to be generated of these people. --- # \n\n')
    empty_entries, df = remove_empty_entries(df)

absent_participants = check_notpresent(df)


import numpy as np
import sys

def remove_in_activities(df):
    df = df[~df['Program'].isin(Activities)]
    invalid_programs = df[~df['Program'].isin(Houses + Activities)]
    if not invalid_programs.empty:
        print('# --- (ERROR) invalid programs: --- #')
        print(invalid_programs)
        print(invalid_programs['Program'])
        raise ValueError("Error: Some entries in 'Program' are neither in Houses nor Activities. Invalid entries:")

    return df

def match_databse_namesplit(df_data):
    df_data_np = df_data.to_numpy()
    new_matched_np = []
    new_no_matched_np = []
    for part in absent_participants['Participant external ID']:
        part = str(int(part))

        foundit = False
        for part2 in df_data_np:
            if str(int(part2[-1])) == part:
                new_matched_np.append(part2)
                foundit = True
                
        if foundit == False:
            print('Error matching participant:')
            print(part)
            sys.exit('ERROR IN MATCHING PARTICIPANTS')
            new_no_matched_np.append(part)

    return np.array(new_matched_np), np.array(new_no_matched_np)
      

def match_databse_nosplit(df_data):
    nonamesplit_df = pd.DataFrame()
    nonamesplit_df = absent_participants['Participant external ID']
    df_data_np = df_data.to_numpy()

    merged_df = pd.merge(nonamesplit_df, df_data, on=['Participant external ID'], how='left', indicator=True)
    # Find rows that are not matched
    unmatched_rows = merged_df[merged_df['_merge'] == 'left_only']
    # Raise an error if there are unmatched rows

    if not unmatched_rows.empty:
        print(unmatched_rows)
        print('unmatched rows')
        error_message = f"Error: The following rows from df1 did not match any person in the database:\n"+ str(unmatched_rows[['Participant  \nFirst name', 'Participant  \nLast name']])
        print(f'\n\n WARNING: {error_message} \n \n')
    #Extract the matched rows from df2
    matched_rows = merged_df[merged_df['_merge'] == 'both']
    matched_rows = matched_rows.drop(columns=['_merge'])

    # Find rows that are not matched
    unmatched_rows = merged_df[merged_df['_merge'] == 'left_only']
    return matched_rows, unmatched_rows


def return_matches_w_database(file_path_data):
    df_data = pd.read_csv(file_path_data, encoding = "utf-8")
    df_data = df_data.map(garbage_remover)

    # No housing database:
    df_data_act = remove_in_house(df_data)
    
    
    testinglukaact = df_data_act.to_numpy()

    df_data_house = remove_in_activities(df_data)
    act_db_match, act_db_nomatch = match_databse_namesplit(df_data_act)
    house_db_match, house_db_nomatch = match_databse_namesplit(df_data_house)

    return house_db_match, act_db_match

def get_house_data_yesterday(file_path):
    df_data = pd.read_csv(file_path, encoding = "utf-8")
    df_data = df_data.map(garbage_remover)

    # Pattern: comma + space + 3-letter month + rest-of-line
    date_pattern = r',\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b.*$'

    df_data['Program'] = (
        df_data['Program']
          .astype(str)                        # ensure string dtype
          .str.replace(date_pattern, '',      # drop “, Jun 2, 2025(…”
                       regex=True)
          .str.strip()                        # clean up any trailing spaces
    )

    df_data = remove_in_activities(df_data)
    
    yesterday_house_match, yesterday_house_nomatch = match_databse_nosplit(df_data)
    return yesterday_house_match
    

    
house_db_match, act_db_match = return_matches_w_database(database_path)
yesterday_house_match = get_house_data_yesterday(yesterday_housing_attendance)


yesterday_extractinfo = ['Participant', 'Date', 'Status', 'Attendance Status', 'Program','Check in', 'Check out']

absent_numpy = absent_participants.to_numpy()

housing_needed_info = house_db_match
act_needed_info = act_db_match
yesterday_needed_info = yesterday_house_match[yesterday_extractinfo].to_numpy()

personlist = []
for person in absent_numpy:

    persondict = {}
    persondict['pacific ID'] = str(int(person[8]))
    persondict['fullname'] = person[0]
    persondict['first name'] = persondict['fullname'].split()[0]
    persondict['last name'] = persondict['fullname'].split()[1:]
    
    # Deal with middle names
    if len(persondict['last name']) > 1:
        newlastname = ''
        for i in persondict['last name']:
            newlastname = newlastname + ' ' + i
        persondict['last name'] = newlastname.lstrip()
    else:
        persondict['last name'] = persondict['last name'][0]
            
    persondict['enroldate'] = person[1]
    persondict['status'] = person[2]
    persondict['attendance status'] = person[3]
    persondict['current activity'] = person[4]
    persondict['checkin'] = person[5]
    persondict['checkout'] = person[6]
    persondict['checkedinby'] = person[7]
    persondict['pickedupby'] = person[8]

    foundhouse = False
    
    for houseinf in housing_needed_info:
        #print(houseinf[20])
        #print(persondict['pacific ID'])
        if str(int(houseinf[20])) == persondict['pacific ID']:
            persondict['house'] = houseinf[3]
            persondict['room'] = houseinf[4]
            foundhouse = True

            
    if foundhouse == False:
        print(f"ERROR \n name: {persondict['fullname']}")
        print('''# ------------------ #
# MAKE THIS AN ERROR 111 #
# ------------------ #''')
        print('ERROR DID NOT FOUND HOUSE CHECK MIDDLE LAST NAME PROBLEM THINGY')

    foundactivity = 0
    for actinf in act_needed_info:

        if str(int(actinf[20])) == persondict['pacific ID']:
            if foundactivity > 0:
                persondict['all programs'] = persondict['all programs'] + ' , ' + actinf[0]
            else:
                persondict['all programs'] = actinf[15]
                persondict['gender'] = actinf[2]
                persondict['date of birth'] = 'nan' #actinf[4]
                persondict['student phone'] = str(int(actinf[5]))
                persondict['P1 firstname'] = actinf[7]
                persondict['P1 lastname'] = actinf[8]
                persondict['P1 rel'] = actinf[10]
                persondict['P1 phone'] = str(int(actinf[9]))
                persondict['P1 authorized'] = 'nan'
                persondict['P2 firstname'] = actinf[11]
                persondict['P2 lastname'] = actinf[12]
                persondict['P2 rel'] = actinf[14]
                persondict['P2 phone'] = str(int(actinf[13]))
                persondict['P2 authorized'] = 'nan'
                persondict['schedule title'] = actinf[16]
                persondict['participant ID'] = str(actinf[19])

                #persondict['pacific ID'] = str(int(actinf[20]))

            foundactivity += 1
    if foundactivity == 0:
        print(f"ERROR \n name: {persondict['fullname']}")
        print('''# ------------------ #
# MAKE THIS AN ERROR 222 #
# ------------------ #''')
        print('ERROR DID NOT FOUND ACTIVITY CHECK MIDDLE LAST NAME PROBLEM THINGY')

    for yesterdayinf in yesterday_needed_info:
        if yesterdayinf[0] == persondict['fullname']:
            persondict['ydate'] = yesterdayinf[1]
            persondict['ystatus'] = yesterdayinf[2]
            persondict['yattendance status'] = yesterdayinf[3]
            persondict['yprogram'] = yesterdayinf[4]
            persondict['ycheckin'] = yesterdayinf[5]
            persondict['ycheckout'] = yesterdayinf[6]
  
    personlist.append(persondict)
    
# Add faculty phone number:
facultyinfo = pd.read_csv(facinfopath, encoding = "utf-8")
facultyinfo = facultyinfo.map(garbage_remover).to_numpy()

for person in personlist:
    for key in person:
        person[key] = str(person[key])
onlynamelist = []
for person in personlist:
    onlynamelist.append(person['fullname'])
    gotfac = 0
    for facprog in facultyinfo:
        prog = facprog[0]
        
        personprog = person['current activity']
        if personprog.strip() == str(prog).strip():
                person['enrProgram'] =  [facprog[0]]
                person['enrfaculty'] = [facprog[1]]
                person['fac contact person'] = [facprog[2]]
                person['fac phonenr'] = [facprog[3]]
                gotfac += 1
                
    if gotfac == 0:
        print('# -- Didn\'t find faculty of: -- #')
        print(person['current activity'])
        print('$ ------------------------------ #')
        person['enrProgram'] =  ['N/A']
        person['enrfaculty'] = ['N/A']
        person['fac contact person'] = ['N/A']
        person['fac phonenr'] = ['N/A']


def get_roommate(file_path_data, personlist):
    df_data = pd.read_csv(file_path_data, encoding = "utf-8")
    df_data = df_data.map(garbage_remover)
    df_data_house = remove_in_activities(df_data).to_numpy()
    df_data_act = remove_in_house(df_data).to_numpy()
    
    for person in personlist:
 
        nrroommates = 0

        room = person['room'][0:-1]
        
        for roommate in df_data_house:

            if type(roommate[4]) is float:
                continue

            roomnrcheck = roommate[4][0:-1]
            if roomnrcheck == room and str(int(roommate[-1])) != person['pacific ID']: #roommate[1].strip() + ' ' + roommate[2].strip() != person['fullname']:
                if nrroommates > 1:
                    print(roommate[1].strip() + ' ' + roommate[2].strip())
                    print(person['fullname'])
                    print(roomnrcheck)
                    print('Person has multiple roommates?? code can\'t deal with this at the moment')
                    sys.exit('Person has multiple roommates?? code can\'t deal with this at the moment')
                # Next feature: Tell when both roommates are not present
                person['rmname'] =  roommate[0].strip() + ' ' + roommate[1].strip()
                person['rmhouse'] = roommate[3]
                person['rmroom'] = roommate[4]

                if roommate[0].strip() + ' ' + roommate[1].strip() in onlynamelist:
                    person['rmabsent'] = 'Absent'
                else:
                    person['rmabsent'] = 'Present or checked out (currently not differentiated)'
                    
                for rminfo in df_data_act:
 
                    if rminfo[0].strip() == roommate[0].strip() and rminfo[1].strip() == roommate[1].strip():

                        person['rmprogram'] = rminfo[15]
                        person['rmphone'] = str(int(rminfo[5]))
                nrroommates += 1
            
            else:
                continue
        if nrroommates == 0:
            print(f"(WARNING) PERSON HAS NO ROOMMATES?!: {person['fullname']}")
    return personlist
            
    
personlist = get_roommate(database_path, personlist)
print('# -- (LOGGER) All data matching SUCCESSFUL for: -- #')
for pi in personlist:
    print(pi['fullname'],  end=', ')
print('\n# ------------------------------------------------ #\n\n')

from docx import Document
from docx.shared import Pt

# Create a new document
doc = Document()

# Add a header
print('# ------------- Generating .docx Document... ------------- #')
print('PLEASE NOTE: When a PERMISSION ERROR arises, you have the document open in a different program and need to close it or change the docxprintdoc variable to a new document ')

for person in personlist:
    # if person['fullname'] in faultypeople:
    #     continue
    # if person['current activity'] == 'Building Electrical Circuits and Computers' or person['current activity'] == 'Creative Writing and Storytelling':
    #     continue
    header = doc.add_heading('Student attendance check data sheet', level=1)
    header.style.font.size = Pt(20)
    # p = doc.add_paragraph()
    # p.add_run('Welcome to the hunt, your job is to track down and demolish following target \n').bold = True
    p = doc.add_paragraph()
    p.add_run('Name: ').bold = True
    p.add_run(person['fullname'] + '\t' + '\t')
    #print(person)
    p.add_run('gender: ').bold = True
    p.add_run(person['gender'] + '\n')  
    # p.add_run('date of birth: ').bold = True
    # p.add_run(person['date of birth'] + '\n')  
    p.add_run('Phone number: ').bold = True
    p.add_run(person['student phone'] + '\n') 
    p.add_run('Program: ').bold = True
    p.add_run(person['current activity'] + '\n')
    p.add_run('pacific ID: ').bold = True
    p.add_run(person['pacific ID'] + '\n') 
    #print(person['pacific ID'])
    p.add_run('House: ').bold = True
    p.add_run(person['house'] + '\t' + '\t')
    p.add_run('Room: ').bold = True
    p.add_run(person['room'] + '\n') 
    p.add_run('status: ').bold = True
    p.add_run(person['status'] + '\n')
    p.add_run('attendance status: ').bold = True
    p.add_run(person['attendance status'] + '\n')
    p.add_run('check in at: ').bold = True
    p.add_run(person['checkin'] + '\t' + '\t')
    p.add_run('check out at: ').bold = True
    p.add_run(person['checkout'] + '\n')

    header = doc.add_heading('Information about his/her roommate', level=2)
    header.style.font.size = Pt(12)
    p = doc.add_paragraph()
    if 'rmname' in person.keys():
        p.add_run('name: ').bold = True
        p.add_run(person['rmname']+ '\n')    
        p.add_run('house: ').bold = True
        p.add_run(person['rmhouse']+ '\t'+ '\t')    
        p.add_run('room: ').bold = True
        p.add_run(person['rmroom']+ '\n')    
        p.add_run('attendance status: ').bold = True
        p.add_run(person['rmabsent']+ '\n')
        p.add_run('program: ').bold = True
        if 'rmprogram' in person.keys():
            p.add_run(person['rmprogram']+ '\n')  
            p.add_run('phone number: ').bold = True
            p.add_run(person['rmphone']+ '\n')    

        else:
            p.add_run('HAS NO PROGRAM PROBABLY SCHOOL DISTRICT WAITING LIST \n')

    
    
    header = doc.add_heading('Parent information', level=2)
    header.style.font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run('Parent 1: ').bold = True
    p.add_run(person['P1 firstname']+ ' ')
    p.add_run(person['P1 lastname']+ '\t' + '\t' + '\t' )
    p.add_run('Parent 2: ').bold = True
    p.add_run(person['P2 firstname'] + ' ')    
    p.add_run(person['P2 lastname']+ '\n')
    
    p.add_run('Parent 1 Language:  ').bold = True
    p.add_run(person['P1 rel']+ '\t' + '\t')    
    p.add_run('Parent 2 Language: ').bold = True
    p.add_run(person['P2 rel']+ '\n')   
    
    p.add_run('Parent 1 phone number: ').bold = True
    p.add_run(person['P1 phone']+ '\t')
    p.add_run('Parent 2 phone number: ').bold = True
    p.add_run(person['P2 phone']+ '\n')
    
    header = doc.add_heading('Last night’s housing attendance check', level=2)
    header.style.font.size = Pt(12)
    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run(person['ydate']+ '\n')    
    p.add_run('Status: ').bold = True
    p.add_run(person['ystatus']+ '\n')    
    p.add_run('Attendance status: ').bold = True
    p.add_run(person['yattendance status']+ '\n')    
    p.add_run('Program/House: ').bold = True
    p.add_run(person['yprogram']+ '\n')    
    p.add_run('Check in: ').bold = True
    p.add_run(person['ycheckin']+ '\t')
    p.add_run('Check out: ').bold = True
    p.add_run(person['ycheckout']+ '\n')       
    header = doc.add_heading('Faculty information', level=2)
    header.style.font.size = Pt(12) 
    p = doc.add_paragraph()

    for ind, faculty in enumerate(person['enrProgram']):
        
        # if person['current activity'].strip() != faculty.strip():
        #     continue
        p.add_run('Program: ').bold = True
        p.add_run(person['enrProgram'][ind] + '\n')    
        #p.add_run('Faculty: ').bold = True
        #p.add_run(person['enrfaculty'][ind] + '\n')    
        p.add_run('Contact person: ').bold = True
        p.add_run(person['fac contact person'][ind] + '\n')    
        p.add_run('Contact person phone number: ').bold = True
        p.add_run(str(person['fac phonenr'][ind]).strip('?') + '\n')   
        p = doc.add_paragraph()
        
    p = doc.add_page_break()

doc.save(docxprintdoc)
print('# -------- DONE Generating the attendance report --------- #')

